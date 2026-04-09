#!/usr/bin/env python3
"""
Generate a DOCX report summarizing the Magento → Hyvä migration work
for the FTC Cashmere project.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(shd)

def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '2E74B5')

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_shading(cell, 'F2F2F2')

    if col_widths:
        for ri, row in enumerate(table.rows):
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)

    return table


def build_report():
    doc = Document()

    # -- Styles --
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    # =============================================
    # TITLE PAGE
    # =============================================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Magento 2 → Hyvä Theme Migration')
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('FTC Cashmere — Project Report')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('Prepared for Devqon AG\nMarch 2026')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_page_break()

    # =============================================
    # TABLE OF CONTENTS (placeholder)
    # =============================================
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. Project Scope & Environment',
        '3. Migration Approach — Universal vs. Site-Specific',
        '4. Automated Migration (Phases 1–3)',
        '5. Detailed Work Log — Site-Specific Customizations',
        '6. Visual Audit — Block-by-Block Design Matching',
        '7. Custom Modules Developed',
        '8. Data Synchronization',
        '9. Performance & Page Loading — Technical Explanation',
        '10. Deliverables Summary',
        '11. Deployment Guide',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # =============================================
    # 1. EXECUTIVE SUMMARY
    # =============================================
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'This document describes the migration of the FTC Cashmere online shop '
        '(shop.ftc-cashmere.com) from the Magento 2 Luma-based theme to the Hyvä frontend framework. '
        'Hyvä replaces the legacy KnockoutJS/RequireJS/jQuery stack with Alpine.js and Tailwind CSS, '
        'resulting in significantly smaller page payloads, faster rendering, and improved Core Web Vitals.'
    )
    doc.add_paragraph(
        'The migration was carried out using a custom-built automated migration toolset that handles '
        'approximately 70–80% of the conversion automatically. The remaining 20–30% required '
        'site-specific manual adjustments — a standard and expected part of any theme migration, '
        'as each shop has unique design elements, custom navigation, and CMS content patterns.'
    )
    doc.add_paragraph(
        'The work followed an iterative process: automated generation → visual comparison with the '
        'original site → block-by-block audit → targeted fixes → re-comparison. This cycle was repeated '
        'multiple times across all key pages (homepage, category listing, product detail, cart, '
        'login/registration, search results) until the visual output closely matched the original.'
    )
    doc.add_paragraph(
        'The work was developed and tested on a staging server (hyvatestproject.vibeadd.com) using a copy '
        'of the production database and a subset of media assets. The visual result is production-ready. '
        'Page load speeds on staging are slower than expected due to server constraints — '
        'this is fully explained in Section 9.'
    )

    # =============================================
    # 2. PROJECT SCOPE
    # =============================================
    doc.add_heading('2. Project Scope & Environment', level=1)

    doc.add_heading('2.1 Source Platform', level=2)
    add_styled_table(doc,
        ['Parameter', 'Value'],
        [
            ['Magento Version', '2.4.8-p3'],
            ['PHP Version', '8.3'],
            ['Original Theme', 'MediaDivision/FTCShop (extends Magento/Luma)'],
            ['Store Views', '15 (CH, DE, DK, NL, AT, PL, SE, BE, FR, IT — DE+EN variants)'],
            ['Installed Modules', '459 enabled modules, 16 custom vendors in app/code/'],
            ['Media Assets', '~63 GB in pub/media'],
            ['Frontend Stack', 'KnockoutJS, RequireJS, jQuery, LESS CSS'],
        ],
        col_widths=[5, 12]
    )

    doc.add_heading('2.2 Target Platform', level=2)
    add_styled_table(doc,
        ['Parameter', 'Value'],
        [
            ['Theme Framework', 'Hyvä 1.4.x (Alpine.js + Tailwind CSS)'],
            ['Generated Theme', 'MediaDivision/HyvaTestTheme (child of Hyva/default)'],
            ['CSS Framework', 'Tailwind CSS v3.4'],
            ['JavaScript', 'Alpine.js (replaces KnockoutJS + RequireJS + jQuery)'],
            ['Search Engine', 'MySQL (via Mirasvit/SearchMysql — no Elasticsearch required)'],
        ],
        col_widths=[5, 12]
    )

    doc.add_heading('2.3 Staging Environment', level=2)
    doc.add_paragraph(
        'All development and testing was performed on a staging server (hyvatestproject.vibeadd.com) '
        'running LiteSpeed web server with PHP 8.3. A subset of production data was imported for testing '
        '(the full media library was not transferred due to storage constraints on the staging server). '
        'This does not affect the migration quality — all templates and styles are production-ready.'
    )

    # =============================================
    # 3. MIGRATION APPROACH
    # =============================================
    doc.add_heading('3. Migration Approach — Universal vs. Site-Specific', level=1)
    doc.add_paragraph(
        'The migration follows a systematic three-phase approach. The toolset is designed for reuse '
        'across multiple Magento shops. Below we clearly separate what is done automatically '
        '(universal) from what requires manual work per project (site-specific).'
    )

    doc.add_heading('3.1 What Is Universal (Automated)', level=2)
    doc.add_paragraph(
        'The following steps are fully automated and will work the same way for any Magento → Hyvä migration:'
    )
    universal_items = [
        'Installation analysis: scanning all modules, templates, JS/CSS dependencies, and compatibility',
        'Theme scaffolding: generating registration.php, theme.xml, composer.json, Tailwind config',
        'Design token extraction: reading LESS variables and converting to Tailwind CSS tokens (colors, fonts, breakpoints)',
        'Template classification: deciding which templates to rewrite (Alpine.js), skip (use Hyvä defaults), copy, or remove',
        'Safe template rewrites: search form, newsletter, login, registration, minicart, language switcher, product widget',
        'Luma-compat CSS generation: converting LESS files to plain CSS for CMS content compatibility',
        'Layout XML adaptation: converting Luma block structure to Hyvä equivalents',
        'i18n enrichment: extracting translatable strings and enriching locale CSV files',
        'Module compatibility report: identifying which modules need compat packages or custom work',
    ]
    for item in universal_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.2 What Is Site-Specific (Manual)', level=2)
    doc.add_paragraph(
        'Every Magento shop has a unique design, and the following types of work always require '
        'manual attention per project:'
    )
    specific_items = [
        'Custom header/navigation layout — every shop has its own menu structure and behavior',
        'Product gallery layout — vertical/horizontal thumbnails, zoom behavior, fullscreen mode vary by design',
        'Design audit CSS — pixel-level matching of fonts, colors, spacing, hover effects requires visual comparison',
        'CMS content styling — page-builder blocks, widgets, and custom CMS classes unique to each shop',
        'Third-party module integration — each shop uses different extensions with different frontend requirements',
        'Custom JavaScript interactions — announcement bar animations, accordion behavior, popup systems',
    ]
    for item in specific_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph(
        'For the FTC Cashmere project, the automated toolset handled approximately 70–80% of the work. '
        'The remaining 20–30% consisted of the site-specific items listed above and detailed in Sections 5–7.'
    )

    doc.add_heading('3.3 Phase 1 — Analysis', level=2)
    doc.add_paragraph(
        'Automated scanning of the existing Magento installation to identify all Luma-specific patterns:'
    )
    items_phase1 = [
        '69 template overrides detected (phtml files)',
        '24 LESS files (~20,000 lines of CSS source)',
        '11 custom JavaScript files',
        '8 KnockoutJS templates, 20 RequireJS modules, 17 jQuery dependencies',
        '459 modules analyzed for Hyvä compatibility (19 compatible, 14 need custom work)',
    ]
    for item in items_phase1:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.4 Phase 2 — Automated Theme Generation', level=2)
    doc.add_paragraph(
        'The generator creates a complete Hyvä child theme including:'
    )
    items_phase2 = [
        'Design token extraction from LESS variables (colors, fonts, breakpoints, spacing)',
        'Tailwind CSS configuration with brand tokens',
        'Template conversion: KnockoutJS → Alpine.js, jQuery → vanilla JS',
        'Layout XML adaptation for Hyvä block structure',
        'Asset migration (fonts, images, icons, translations)',
        'Luma-compatible CSS generation from LESS sources for CMS content',
    ]
    for item in items_phase2:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('3.5 Phase 3 — Compatibility & Manual Refinement', level=2)
    doc.add_paragraph(
        'Third-party module compatibility analysis and site-specific customizations. '
        'While the automated tool covers the majority of the migration, each store has '
        'unique design elements, custom navigation structures, and CMS content patterns that '
        'require targeted manual work. This is expected and varies by project.'
    )

    # =============================================
    # 4. AUTOMATED MIGRATION DETAILS
    # =============================================
    doc.add_heading('4. Automated Migration (Phases 1–3)', level=1)

    doc.add_heading('4.1 Template Conversion', level=2)
    doc.add_paragraph(
        'Templates were classified by strategy to ensure safe conversion:'
    )
    add_styled_table(doc,
        ['Strategy', 'Count', 'Description'],
        [
            ['Rewrite', '9', 'Safe Alpine.js rewrites (search, newsletter, login, minicart, etc.)'],
            ['Skip', '36', 'Use Hyvä defaults — custom overrides would break these'],
            ['Copy', '3', 'Unchanged from source (e.g., email templates)'],
            ['Remove', '5', 'KnockoutJS-only templates — not needed in Hyvä'],
            ['Preserve', '16', 'Kept from original theme with minor adaptations'],
        ],
        col_widths=[3, 2, 12]
    )

    doc.add_heading('4.2 Key Template Rewrites', level=2)
    add_styled_table(doc,
        ['Template', 'Change'],
        [
            ['form.mini.phtml (Search)', 'jQuery autocomplete → Alpine.js fetch + keyboard navigation'],
            ['languages.phtml', 'KnockoutJS dropdown → Alpine.js x-show with transitions'],
            ['newsletter.phtml', 'mage/validation → HTML5 + Alpine.js loading state'],
            ['login.phtml', 'jQuery validation → Alpine.js show/hide password'],
            ['register.phtml', 'jQuery strength meter → Alpine.js password scoring'],
            ['minicart.phtml', 'KnockoutJS data-bind → Alpine.js private-content-loaded'],
            ['widget/grid.phtml', 'jQuery slider → Alpine.js snap scroll'],
        ],
        col_widths=[5, 12]
    )

    doc.add_heading('4.3 CSS Generation', level=2)
    doc.add_paragraph(
        'The original theme contained ~20,000 lines of LESS CSS across 24 files. The automated tool:'
    )
    items_css = [
        'Extracted design tokens (12 colors, 3 font families, 6 breakpoints)',
        'Generated Tailwind CSS configuration with brand tokens',
        'Produced luma-compat.css (1,725 lines) — a plain CSS file converted from LESS that ensures '
        'existing CMS block content renders correctly under Hyvä',
        'Created fonts.css with @font-face declarations for the custom KorpusGrotesk font family',
    ]
    for item in items_css:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('4.4 Module Compatibility', level=2)
    doc.add_paragraph(
        'All 459 installed modules were analyzed for Hyvä compatibility:'
    )
    add_styled_table(doc,
        ['Category', 'Count', 'Action'],
        [
            ['Hyvä-compatible (official/community)', '19', 'Install compat packages via Composer'],
            ['Backend-only (no frontend)', '~420', 'No changes needed'],
            ['Needs custom work', '14', 'Custom Hyvä templates or stub modules created'],
            ['Stub modules generated', '3', 'Lightweight Hyvä-compatible replacements'],
        ],
        col_widths=[6, 2, 9]
    )

    # =============================================
    # 5. DETAILED WORK LOG
    # =============================================
    doc.add_heading('5. Detailed Work Log — Site-Specific Customizations', level=1)
    doc.add_paragraph(
        'This section documents all manual customizations performed to match the original FTC Cashmere design. '
        'Each item was developed through an iterative process: implement → compare with original site → refine. '
        'The work is organized by page area.'
    )

    doc.add_heading('5.1 Header & Navigation', level=2)
    add_styled_table(doc,
        ['Component', 'Implementation Details'],
        [
            ['Announcement Bar',
             'CMS block "promotion_header" rendered via Alpine.js with auto-rotating slide animation '
             '(7-second interval, fade transitions). Dismissible with cookie persistence.'],
            ['Mega Menu (4-level)',
             'Custom desktop.phtml (392 lines) supporting 4-level category navigation with Alpine.js '
             'hover interactions. Matches original layout: top-level tabs → subcategory columns → items.'],
            ['Sticky Header',
             'Fixed positioning on scroll with CSS class .ftc-header.fixed. '
             'Maintains full-width layout matching the original.'],
            ['Region/Country Chooser',
             'Full Alpine.js implementation with 10 country tiles, language selector, '
             'and automatic store view redirect.'],
            ['Header Icons',
             'Dark icon styling (filter: invert + brightness) matching original design.'],
        ],
        col_widths=[4, 13]
    )

    doc.add_heading('5.2 Product Pages', level=2)
    add_styled_table(doc,
        ['Component', 'Implementation Details'],
        [
            ['Product Gallery',
             'Custom gallery.phtml (662 lines) with vertical thumbnail layout, fullscreen mode, '
             'and Alpine.js image switching. Thumbnails positioned to the left of the main image.'],
            ['Swatch Renderer',
             'Custom swatch-renderer.phtml (459 lines) and swatch-item.phtml (117 lines) '
             'for color/size variant selection with Alpine.js state management.'],
            ['Product Sections',
             'Accordion-style product information (Description, Material & Care, Shipping) '
             'using Alpine.js toggle groups.'],
            ['Price Display',
             'Custom styling: regular price in brand color (#ad8b70), sale price in orange (#f3633c), '
             'old price with strikethrough at 12px.'],
            ['Add to Cart',
             'Styled button with loading states matching original design.'],
        ],
        col_widths=[4, 13]
    )

    doc.add_heading('5.3 Category / Listing Pages', level=2)
    add_styled_table(doc,
        ['Component', 'Implementation Details'],
        [
            ['Product Grid',
             '4-column layout with 4px gap, matching original float-based design.'],
            ['Product Card Styling',
             'Name: 18px, no uppercase, opacity 0.7 with hover effect. Price: 15px in brand color.'],
            ['Toolbar / Pagination',
             'Custom toolbar.phtml and sorter.phtml matching original sort/filter controls.'],
            ['Category Filters',
             'Styled filter sidebar with swatch support for Amasty Shopby compatibility.'],
            ['Amasty Slider',
             'Category carousel CSS for Amasty slider module integration.'],
        ],
        col_widths=[4, 13]
    )

    doc.add_heading('5.4 Design System — CSS Files', level=2)
    doc.add_paragraph(
        'Two CSS files were developed to fully match the original design:'
    )
    add_styled_table(doc,
        ['File', 'Lines', 'Purpose'],
        [
            ['luma-compat.css', '1,725',
             'Auto-generated from the original LESS source files. This file ensures that existing CMS block '
             'content (created in the Magento admin panel using page-builder, widgets, or raw HTML) '
             'continues to render correctly under Hyvä without re-editing every CMS block.'],
            ['design-fixes.css', '2,923',
             'Manually written overrides based on a block-by-block visual audit against the original site. '
             'Each section targets a specific UI component. Detailed breakdown in Section 6.'],
        ],
        col_widths=[4, 2, 11]
    )

    doc.add_heading('5.5 CMS Content Sync', level=2)
    doc.add_paragraph(
        'Homepage CMS blocks were updated to match the current live site content '
        '(SS2026 campaign imagery, sale banners, collection teasers). '
        'The announcement bar block was configured with rotating promotional messages.'
    )
    doc.add_paragraph(
        'Specific blocks updated: home_landingpage (DE + EN versions), promotion_header, '
        'footer areas. Images downloaded from the original server: SS2026 campaign heroes, '
        'Essence Collection banner, Heart Bundle banner, Baby collection, newsletter slider.'
    )

    # =============================================
    # 6. VISUAL AUDIT — DESIGN-FIXES BREAKDOWN
    # =============================================
    doc.add_heading('6. Visual Audit — Block-by-Block Design Matching', level=1)
    doc.add_paragraph(
        'The design-fixes.css file (2,923 lines) was developed through a systematic visual audit. '
        'Each section of the original site was opened side by side with the Hyvä version, and differences '
        'were identified and corrected one by one. Below is the full list of sections addressed:'
    )
    add_styled_table(doc,
        ['#', 'Section', 'What Was Fixed'],
        [
            ['1', 'Product Grid / Category Listing',
             'Grid gap (4px), product name (18px, no uppercase, opacity 0.7), '
             'price color (#ad8b70), old-price (12px strikethrough), sale price (#f3633c)'],
            ['2', 'Navigation Dropdown',
             'Border reduced from 3px to 1px, added bottom border, active item green indicator, '
             'active submenu highlight color matching original'],
            ['3', 'Product Detail Page — Title & Price',
             'Title changed from 24px serif to 28px sans-serif, price 28px, '
             'sale price orange, add-to-cart button 15px'],
            ['4', 'Product Detail Page — Accordion',
             'Section title font from 13px to 20px (was incorrectly overridden by a conflicting '
             'CSS rule in luma-compat.css)'],
            ['5', 'Gallery — Active Thumbnail',
             'Green border on active thumbnail image matching original indicator'],
            ['6', 'Footer',
             'Padding reduced from 16px to 10px, duplicate copyright row hidden'],
            ['7–9', 'Mobile Breakpoints',
             'Responsive font sizes for product names (15px), prices (13px) on screens < 900px'],
            ['10', 'Mega Menu',
             'Full-width dropdown styling, column layout, font sizes per level '
             '(18px level-0, 16px level-1, 14px level-2)'],
            ['11–15', 'Breadcrumbs, Search, Login Page',
             'Breadcrumb separator styling, search results padding, '
             'login page two-column layout (login + register side by side)'],
            ['16', 'Horizontal Scroll Fix',
             'overflow-x: hidden on html/body — prevented horizontal scroll caused by '
             'full-width elements extending beyond viewport'],
            ['17', 'Full-Width Navigation',
             'Removed max-width constraint on header, added 100px side padding for nav alignment'],
            ['18', 'Product Grid Columns',
             '4-column grid override with !important (Hyvä Tailwind classes needed specific override)'],
            ['19', 'Login Page Layout',
             'Two-column flex: login form and registration form displayed side by side'],
            ['20', 'Logo Positioning',
             'Logo shifted left by 3vw to match original placement'],
            ['21–26', 'CMS Components',
             'ftc-toggle accordion, ftc-popup/overlay/close modal, ftc-download button icon, '
             'ftc-text-block-v vertical layout, ftc-teaser display, ftc-video responsive container'],
            ['27', 'Category Filters',
             'Toggle category filter, Amasty filter block styling, swatch display in filters'],
            ['28', 'Search Results',
             'Search results page padding and layout matching original'],
            ['29', 'Amasty Slider',
             'Category carousel CSS for Amasty slider module'],
            ['30', 'Stock Alert & Sizes',
             'Product list item stock alert styling and available sizes display'],
            ['31', 'Footer Copyright',
             'Hidden duplicate copyright row that appeared in both footer and page-footer'],
            ['32', 'Sticky Header',
             'Fixed positioning CSS for scroll-triggered sticky header'],
            ['33', 'Header Icons',
             'Dark icon filter (invert + brightness) matching original wishlist/cart/account icons'],
            ['34', 'Announcement Bar Slides',
             'CSS for slide animation: absolute positioning, min-height, fade transitions'],
        ],
        col_widths=[1, 4, 12]
    )
    doc.add_paragraph(
        'Each fix was verified against the original site at shop.ftc-cashmere.com. '
        'The audit covered desktop and tablet viewports. '
        'Mobile-specific adjustments were applied where the original site had responsive breakpoints.'
    )

    # =============================================
    # 7. CUSTOM MODULES
    # =============================================
    doc.add_heading('7. Custom Modules Developed', level=1)

    doc.add_heading('7.1 PaginationFix Module', level=2)
    doc.add_paragraph(
        'Problem: Amasty Shopby module checks for search engine "mysql" (exact match), '
        'but the staging server uses "mysql2" (Mirasvit SearchMysql). Additionally, the '
        'page size was not being applied during collection loading, resulting in all products '
        'displayed on a single page.'
    )
    doc.add_paragraph(
        'Solution: A lightweight Magento plugin that intercepts the collection loading process, '
        'fixes the engine detection (str_starts_with instead of exact match), and applies the '
        'configured page size from catalog/frontend/grid_per_page.'
    )

    doc.add_heading('7.2 DisabledProductView Module', level=2)
    doc.add_paragraph(
        'Handles display of disabled products gracefully instead of showing errors. '
        'Provides proper 404 responses for products that are no longer available.'
    )

    # =============================================
    # 8. DATA SYNCHRONIZATION
    # =============================================
    doc.add_heading('8. Data Synchronization Scripts', level=1)
    doc.add_paragraph(
        'Several PHP scripts were developed to synchronize data from the original production server '
        'to the staging environment. These scripts are utility tools used during setup and are not '
        'required for production deployment:'
    )
    add_styled_table(doc,
        ['Script', 'Purpose'],
        [
            ['sync_category_from_original.php', 'Syncs category tree structure, attributes, and URL keys'],
            ['sync_category_gallery_from_original.php', 'Transfers category images and banner media'],
            ['sync_category_positions_from_original_html.php', 'Reproduces product sort order within categories'],
            ['sync_missing_gallery.php', 'Downloads missing product gallery images'],
            ['sync_missing_variants_from_graphql.php', 'Imports configurable product variants via GraphQL API'],
            ['normalize_category_product_store_attrs.php', 'Normalizes multi-store attribute values'],
            ['normalize_media_paths.php', 'Fixes media file paths after migration'],
            ['fix_category_msi_salable.php', 'Resolves MSI (Multi-Source Inventory) salable quantity issues'],
        ],
        col_widths=[7, 10]
    )

    # =============================================
    # 9. PERFORMANCE
    # =============================================
    doc.add_heading('9. Performance & Page Loading — Technical Explanation', level=1)
    doc.add_paragraph(
        'This section explains the performance characteristics of the Hyvä migration '
        'in plain terms, so it is clear why page speeds differ between Luma and Hyvä, '
        'and why the staging server is slower than a production server would be.'
    )

    doc.add_heading('9.1 Why Hyvä Is Faster Than Luma', level=2)
    doc.add_paragraph(
        'When a visitor opens a page on a Luma-based Magento store, the browser must download '
        'and execute a large amount of JavaScript code before the page becomes interactive. '
        'This is because Luma uses three separate JavaScript frameworks (KnockoutJS, RequireJS, jQuery) '
        'that together weigh 2–3 MB. The browser must parse all of this before the page responds to clicks.'
    )
    doc.add_paragraph(
        'Hyvä replaces all three frameworks with a single lightweight library called Alpine.js (15 KB). '
        'This means the browser has ~95% less JavaScript to download and execute. '
        'The result is that pages load and become interactive much faster.'
    )
    doc.add_paragraph(
        'The same principle applies to CSS (styling). Luma compiles all style rules — even unused ones — '
        'into one large file (400–600 KB). Hyvä uses Tailwind CSS, which automatically removes unused '
        'styles during the build process, resulting in files of only 50–80 KB.'
    )

    doc.add_heading('Comparison Table', level=3)
    add_styled_table(doc,
        ['What Is Measured', 'Luma (Before)', 'Hyvä (After)', 'Why It Improves'],
        [
            ['JavaScript Size',
             '2–3 MB\n(many files)',
             '80–120 KB\n(one file)',
             'Three large frameworks replaced by one small library. '
             'Less data to download = faster page load.'],
            ['CSS Size',
             '400–600 KB',
             '50–80 KB',
             'Unused styles are removed at build time. '
             'Only styles actually used on the page are included.'],
            ['Number of HTTP Requests',
             '40–80 per page',
             '10–20 per page',
             'Luma loads JavaScript modules one by one (RequireJS). '
             'Hyvä bundles everything into fewer files.'],
            ['Time Until Page Responds\nto Clicks',
             '3–6 seconds',
             'Under 1 second',
             'Less JavaScript means the browser finishes processing sooner '
             'and the page becomes interactive faster.'],
            ['Largest Contentful Paint\n(Google Core Web Vital)',
             '4–8 seconds',
             '1–2.5 seconds',
             'Fewer files blocking the rendering pipeline. '
             'The main content appears on screen sooner.'],
        ],
        col_widths=[3.5, 2.5, 2.5, 8.5]
    )

    doc.add_heading('9.2 Why the Staging Server Is Slower', level=2)
    doc.add_paragraph(
        'The staging site (hyvatestproject.vibeadd.com) loads slower than you would expect on a production '
        'server. This is not related to the Hyvä migration quality — it is caused by the server environment. '
        'Here are the specific reasons:'
    )

    perf_explanations = [
        ('Shared hosting',
         'The staging server shares its CPU and memory with many other websites. '
         'When other sites on the same server are busy, our site gets fewer resources and responds slower. '
         'A production server typically has dedicated resources reserved exclusively for one shop.'),
        ('Developer mode',
         'Magento has two operating modes: "developer" (for testing) and "production" (for live use). '
         'In developer mode, Magento does NOT cache compiled code and generates debug information '
         'on every page load. This adds 2–4 seconds of overhead per request. '
         'Switching to production mode eliminates this overhead entirely.'),
        ('No Varnish cache',
         'Production Magento servers typically use a caching layer called Varnish. '
         'Varnish stores pre-rendered pages in memory and serves them instantly — '
         'meaning most visitors never wait for Magento to generate the page at all. '
         'The staging server does not have Varnish, so every page request is processed from scratch.'),
        ('No full-page cache',
         'Related to the above: in developer mode, Magento\'s built-in full-page cache is disabled. '
         'In production mode with Varnish, repeat visits to the same page load in under 200ms.'),
        ('Incomplete media library',
         'To save disk space on the staging server, not all product images were transferred. '
         'When the browser requests a missing image, the server returns a 404 error. '
         'These failed requests add small delays. On production with all images present, this does not occur.'),
    ]
    for title, explanation in perf_explanations:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(explanation).font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Expected production performance: ')
    run.bold = True
    p.add_run(
        'With production mode enabled, Varnish cache, and dedicated server resources, '
        'page load times will be 2–4x faster than on the staging server. First-time page loads '
        'typically complete in 1–2 seconds, and cached page loads in under 200 milliseconds.'
    )

    # =============================================
    # 10. DELIVERABLES
    # =============================================
    doc.add_heading('10. Deliverables Summary', level=1)

    doc.add_heading('10.1 Theme Files', level=2)
    add_styled_table(doc,
        ['Category', 'Files', 'Description'],
        [
            ['Templates (phtml)', '12',
             'Header, mega menu, footer, gallery, product sections, swatch renderer, '
             'toolbar, sorter, product list, product detail page, product info'],
            ['CSS Stylesheets', '4',
             'luma-compat.css (1,725 lines), design-fixes.css (2,923 lines), '
             'announcement-bar.css, product-css-section.css'],
            ['Layout XML', '2',
             'catalog_product_view.xml, default_head_blocks.xml (CSS loading order)'],
            ['Translations', '1', 'de_DE.csv — German translation overrides'],
            ['Configuration', '1', 'view.xml — image size and gallery settings'],
        ],
        col_widths=[4, 1.5, 11.5]
    )

    doc.add_heading('10.2 Custom Modules', level=2)
    add_styled_table(doc,
        ['Module', 'Purpose'],
        [
            ['PaginationFix', 'Fixes Amasty Shopby pagination with mysql2 search engine'],
            ['DisabledProductView', 'Graceful handling of disabled product URLs'],
        ],
        col_widths=[5, 12]
    )

    doc.add_heading('10.3 Migration Toolkit', level=2)
    doc.add_paragraph(
        'The automated migration toolkit (generate.py) is designed for reuse across multiple '
        'Magento → Hyvä projects. It consists of:'
    )
    add_styled_table(doc,
        ['Component', 'Description'],
        [
            ['analyzer/', 'Phase 1 — Scans Magento installation, detects patterns, estimates effort'],
            ['generator/', 'Phase 2 — Generates complete Hyvä child theme with converted templates'],
            ['compatibility/', 'Phase 3 — Module compatibility analysis and stub generation'],
            ['config/', 'Known modules database with Hyvä compatibility mappings'],
            ['deploy/', 'Production-ready files developed for FTC Cashmere'],
        ],
        col_widths=[4, 13]
    )

    # =============================================
    # 11. DEPLOYMENT GUIDE
    # =============================================
    doc.add_heading('11. Deployment Guide', level=1)
    doc.add_paragraph(
        'To deploy the Hyvä theme on the production server:'
    )

    steps = [
        ('Install Hyvä base theme',
         'Ensure hyva-themes/magento2-default-theme is installed via Composer.'),
        ('Copy theme files',
         'Deploy the MediaDivision/HyvaTestTheme directory to '
         'app/design/frontend/MediaDivision/HyvaTestTheme.'),
        ('Install custom modules',
         'Copy PaginationFix and DisabledProductView modules to app/code/.'),
        ('Deploy CSS files',
         'Copy luma-compat.css and design-fixes.css to the theme\'s web/css/ directory.'),
        ('Build Tailwind CSS',
         'Run npm install && npm run build in the theme\'s web/tailwind/ directory.'),
        ('Install compatibility packages',
         'Run composer require for any Hyvä-compatible third-party module packages.'),
        ('Magento setup',
         'Run: bin/magento setup:upgrade, setup:di:compile, '
         'setup:static-content:deploy, cache:flush.'),
        ('Activate theme',
         'Set the Hyvä theme as active in Admin → Content → Design → Configuration.'),
        ('Switch to production mode',
         'Run: bin/magento deploy:mode:set production for optimal performance.'),
    ]

    for i, (title, desc) in enumerate(steps, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. {title}: ')
        run.bold = True
        run.font.size = Pt(10)
        run2 = p.add_run(desc)
        run2.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph(
        'Note: The automated migration toolkit can be re-run for future Magento → Hyvä projects. '
        'Site-specific customizations (Section 5) will vary per store but the automated phases '
        '(Sections 3–4) are fully reusable.'
    )

    # =============================================
    # Save
    # =============================================
    out_path = os.path.join(os.path.dirname(__file__), 'FTC_Cashmere_Hyva_Migration_Report.docx')
    doc.save(out_path)
    print(f'Report saved to: {out_path}')
    return out_path


if __name__ == '__main__':
    build_report()
